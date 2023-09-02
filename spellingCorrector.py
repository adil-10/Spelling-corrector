#itterate over all word documents in a folder, if any spelling mistakes are found, fix then with the most suggested one

import os
import docx
from autocorrect import Speller

#specific to .docx exntension
ext = ('.docx')
spell = Speller()

#itterate over Test folder
for files in os.listdir("/Users/adilbadat/Documents/projects/practicePython/Test"):
    #ensure only files with .docx entensions are found
    if files.endswith(ext):
        #printing out ontents of each .docx file located within the Test folder
        full_path = os.path.join("/Users/adilbadat/Documents/projects/practicePython/Test", files)
        doc = docx.Document(full_path)
        result = [p.text for p in doc.paragraphs]
        for w in doc.paragraphs:
            # Fix spelling
            object_strng = str(result) 
            finalSpelling = spell(object_strng)
            # Clear existing content
            clearPara = w.clear()
             # Add the corrected text to the paragraph
            w.add_run(finalSpelling)
        # Save the modified document
        doc.save(full_path)

        


      
   
   